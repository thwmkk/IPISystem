"""Генерация отчёта в формате Word (python-docx)."""
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ACCENT = RGBColor(0x1E, 0x3A, 0x5F)
NEAR_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
DARK_GRAY = RGBColor(0x55, 0x55, 0x55)


def _set_cell_shading(cell, color_hex):
    """Установить фон ячейки таблицы."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def _add_paragraph(doc, text, size=11, bold=False, color=NEAR_BLACK, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    return p


def _add_heading(doc, text, size=14, color=ACCENT):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = color
    return p


def _add_table(doc, headers, rows, col_widths_cm=None):
    """Добавить таблицу с заголовком и данными."""
    if not rows and not headers:
        return None
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'

    # Заголовок
    hdr_row = table.rows[0]
    for ci, h in enumerate(headers):
        cell = hdr_row.cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, '1E3A5F')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Данные
    for ri, row_data in enumerate(rows, start=1):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 0:
                _set_cell_shading(cell, 'F5F8FC')

    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[ci].width = Cm(w)

    return table


def _add_employee_section(doc, emp_data):
    """Добавить секцию отчёта по сотруднику."""
    profile = emp_data['employee']
    _add_heading(doc, profile['full_name'], size=18)

    # Профиль
    _add_heading(doc, 'Профиль сотрудника', size=13)
    profile_table = doc.add_table(rows=6, cols=2)
    profile_table.style = 'Light List'
    rows_data = [
        ('Должность', profile['position']),
        ('Подразделение', profile['department']),
        ('Возраст', str(profile['age'])),
        ('Стаж научной работы', f"{profile['experience']} лет"),
        ('Учёная степень', profile['academic_degree']),
        ('Email', profile['email']),
    ]
    for ri, (k, v) in enumerate(rows_data):
        cell_k = profile_table.rows[ri].cells[0]
        cell_v = profile_table.rows[ri].cells[1]
        cell_k.text = ''
        cell_v.text = ''
        p_k = cell_k.paragraphs[0]; r_k = p_k.add_run(k)
        r_k.font.bold = True; r_k.font.size = Pt(10); r_k.font.name = 'Calibri'
        p_v = cell_v.paragraphs[0]; r_v = p_v.add_run(v)
        r_v.font.size = Pt(10); r_v.font.name = 'Calibri'
        cell_k.width = Cm(5); cell_v.width = Cm(12)
    doc.add_paragraph()

    # IPI
    _add_heading(doc, 'Индивидуальный показатель деятельности (IPI)', size=13)
    ipi = emp_data['ipi']
    ipi_table = doc.add_table(rows=7, cols=2)
    ipi_table.style = 'Light List'
    ipi_rows = [
        ('Итоговый IPI', round(ipi['total'], 2)),
        ('Научная составляющая', round(ipi['scientific'], 2)),
        ('Организационная составляющая', round(ipi['organizational'], 2)),
        ('Техническая составляющая', round(ipi['technical'], 2)),
        ('Коэффициент возраста', ipi['factors']['k_age']),
        ('Коэффициент стажа', ipi['factors']['k_exp']),
        ('Коэффициент аспирантуры', ipi['factors']['k_phd']),
    ]
    for ri, (k, v) in enumerate(ipi_rows):
        cell_k = ipi_table.rows[ri].cells[0]
        cell_v = ipi_table.rows[ri].cells[1]
        cell_k.text = ''
        cell_v.text = ''
        p_k = cell_k.paragraphs[0]; r_k = p_k.add_run(k)
        r_k.font.bold = True; r_k.font.size = Pt(10); r_k.font.name = 'Calibri'
        p_v = cell_v.paragraphs[0]; r_v = p_v.add_run(str(v))
        r_v.font.size = Pt(10); r_v.font.name = 'Calibri'
        cell_k.width = Cm(7); cell_v.width = Cm(5)
    doc.add_paragraph()

    # Научные работы
    _add_heading(doc, f"Научные работы ({len(emp_data['scientific_works'])} шт.)", size=13)
    if emp_data['scientific_works']:
        rows = [
            (w['title'][:80], w['work_type'], str(w['date']),
             str(round(w['points'], 1)), 'Да' if w['verified'] else 'Нет')
            for w in emp_data['scientific_works']
        ]
        _add_table(doc, ['Название', 'Тип', 'Дата', 'Баллы', 'Верифицирована'],
                   rows, col_widths_cm=[7, 4, 2.5, 1.5, 2.5])
    else:
        _add_paragraph(doc, 'Работ за период нет', size=10, color=DARK_GRAY)
    doc.add_paragraph()

    # Организационные работы
    _add_heading(doc, f"Организационные работы ({len(emp_data['organizational_works'])} шт.)", size=13)
    if emp_data['organizational_works']:
        rows = [
            (w['title'][:80], w['work_type'], str(w['date']),
             str(round(w['points'], 1)), 'Да' if w['verified'] else 'Нет')
            for w in emp_data['organizational_works']
        ]
        _add_table(doc, ['Название', 'Тип', 'Дата', 'Баллы', 'Верифицирована'],
                   rows, col_widths_cm=[7, 4, 2.5, 1.5, 2.5])
    else:
        _add_paragraph(doc, 'Работ за период нет', size=10, color=DARK_GRAY)
    doc.add_paragraph()

    # Технические работы
    _add_heading(doc, f"Технические работы ({len(emp_data['technical_works'])} шт.)", size=13)
    if emp_data['technical_works']:
        rows = [
            (w['title'][:80], w['work_type'], str(w['date']),
             str(round(w['points'], 1)), 'Да' if w['verified'] else 'Нет')
            for w in emp_data['technical_works']
        ]
        _add_table(doc, ['Название', 'Тип', 'Дата', 'Баллы', 'Верифицирована'],
                   rows, col_widths_cm=[7, 4, 2.5, 1.5, 2.5])
    else:
        _add_paragraph(doc, 'Работ за период нет', size=10, color=DARK_GRAY)
    doc.add_paragraph()

    # Выполненные задачи
    _add_heading(doc, f"Выполненные задачи ({len(emp_data['tasks_done'])} шт.)", size=13)
    if emp_data['tasks_done']:
        rows = [
            (t['title'][:80], t['project'][:40], t['priority'],
             str(t['deadline']), str(round(t['points'], 1)))
            for t in emp_data['tasks_done']
        ]
        _add_table(doc, ['Название', 'Проект', 'Приоритет', 'Срок', 'Баллы'],
                   rows, col_widths_cm=[6, 5, 2.5, 2, 2])
    else:
        _add_paragraph(doc, 'Выполненных задач за период нет', size=10, color=DARK_GRAY)
    doc.add_paragraph()

    # Невыполненные / просроченные задачи
    _add_heading(doc, f"Невыполненные задачи ({len(emp_data['tasks_pending'])} шт.)", size=13)
    if emp_data['tasks_pending']:
        rows = [
            (t['title'][:80], t['project'][:40], t['priority'],
             t['status'], str(t['deadline']))
            for t in emp_data['tasks_pending']
        ]
        _add_table(doc, ['Название', 'Проект', 'Приоритет', 'Статус', 'Срок'],
                   rows, col_widths_cm=[6, 5, 2.5, 2.5, 2])
    else:
        _add_paragraph(doc, 'Невыполненных задач нет', size=10, color=DARK_GRAY)
    doc.add_paragraph()

    # Проекты
    _add_heading(doc, f"Проекты ({len(emp_data['projects'])} шт.)", size=13)
    if emp_data['projects']:
        rows = [
            (p['name'][:60],
             'Создатель' if p['is_creator'] else 'Участник',
             str(p['start_date']),
             str(p['end_date'] or '—'),
             'Завершён' if p['completed'] else 'В работе')
            for p in emp_data['projects']
        ]
        _add_table(doc, ['Название проекта', 'Роль', 'Начало', 'Окончание', 'Статус'],
                   rows, col_widths_cm=[7, 3, 2.5, 2.5, 2.5])
    else:
        _add_paragraph(doc, 'Проектов нет', size=10, color=DARK_GRAY)


def render_department_word(data) -> bytes:
    """Сгенерировать Word-отчёт по отделу."""
    doc = Document()

    # Поля страницы — пошире
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # Шапка
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Отчёт по деятельности отдела «{data['department']['short_name']}»")
    r.font.name = 'Calibri'; r.font.size = Pt(18); r.font.bold = True
    r.font.color.rgb = ACCENT

    _add_paragraph(doc, f"Период: {data['period']['label']}", size=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"Сформировано: {data['generated_at']}", size=11,
                   color=DARK_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Сводная таблица
    _add_heading(doc, 'Сводная таблица по сотрудникам', size=14)
    headers = ['ФИО', 'Должность', 'IPI', 'Научная', 'Орг.', 'Тех.']
    rows = []
    for emp_data in data['employees']:
        rows.append((
            emp_data['employee']['full_name'],
            emp_data['employee']['position'],
            round(emp_data['ipi']['total'], 2),
            round(emp_data['ipi']['scientific'], 2),
            round(emp_data['ipi']['organizational'], 2),
            round(emp_data['ipi']['technical'], 2),
        ))
    _add_table(doc, headers, rows, col_widths_cm=[5, 4, 1.5, 2, 2, 1.5])

    # Секции по сотрудникам — каждая с новой страницы
    for emp_data in data['employees']:
        doc.add_page_break()
        _add_employee_section(doc, emp_data)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def render_employee_word(emp_data, department_short, period_label, generated_at) -> bytes:
    """Сгенерировать Word-отчёт по одному сотруднику."""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Отчёт по деятельности сотрудника')
    r.font.name = 'Calibri'; r.font.size = Pt(18); r.font.bold = True
    r.font.color.rgb = ACCENT

    _add_paragraph(doc, f"Отдел: {department_short}", size=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"Период: {period_label}", size=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"Сформировано: {generated_at}", size=11,
                   color=DARK_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    _add_employee_section(doc, emp_data)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
